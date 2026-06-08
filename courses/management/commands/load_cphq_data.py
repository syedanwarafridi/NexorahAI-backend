"""
Management command: load_cphq_data

Reads the CPHQ Course folder, parses all Word documents for quiz/mock-exam
questions, copies video files into the Django media directory, and seeds the
database with:

  Category  : Healthcare Quality
  Course    : CPHQ Exam Preparation
  Chapters  : one per Domain (video attached)
  Topics    : one per Domain (Patient Safety, Quality Review & Accountability)
  Questions : parsed from per-domain quiz & mock-exam .docx files, plus the
              global 140-question mock exam

Run with:
    python manage.py load_cphq_data
    python manage.py load_cphq_data --flush   # drop existing data first
"""

import os
import re
import shutil
from pathlib import Path

from django.conf import settings
from django.core.files import File
from django.core.management.base import BaseCommand, CommandError

from courses.models import Category, Chapter, Course
from quiz.models import Question, Topic


# ---------------------------------------------------------------------------
# Docx parsing helpers
# ---------------------------------------------------------------------------

def _extract_text_from_docx(docx_path: Path) -> str:
    """Return all paragraph + table-cell text from a .docx file."""
    try:
        from docx import Document
    except ImportError:
        raise CommandError("python-docx not installed. Run: pip install python-docx")

    doc = Document(str(docx_path))
    lines = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Tables (some docs put Q&A in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)

    return "\n".join(lines)


def _parse_questions(text: str) -> list[dict]:
    """
    Parse MCQ questions from raw text. Handles four formats found in CPHQ docs:

    Format A  — question number on its own line, then text on next line:
        Question 1
        Question text?
        A. Option  ...  ✅ Correct Answer: C  Rationale: ...

    Format B  — question number + text on the SAME line (mock exam docs):
        Question 1 Question text that starts right here...
        A. Option  ...  Correct Answer: B Rationale: ...

    Format C  — Q-prefixed with answer as "✅ B — explanation":
        Q1. Question text?
        A) Option  ...  ✅ B — explanation text

    Format D  — numbered with separator:
        1. Question text?
        A. Option  ...  Answer: A  Explanation: ...
    """
    questions = []
    lines = text.splitlines()

    # Patterns that mark the START of a new question
    q_start_re = re.compile(
        r'^\s*(?:Question\s+(\d+)|Q(\d+)[.)\-:\s]|(\d+)[.):\-]\s)',
        re.IGNORECASE,
    )

    # Collect raw blocks, each starting at a question-start line
    blocks: list[list[str]] = []
    current: list[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        m = q_start_re.match(stripped)
        if m:
            if current:
                blocks.append(current)
            # For "Question N text..." keep the text after the number;
            # for "Question N" alone keep an empty first line (text comes next).
            num_group = m.group(1) or m.group(2) or m.group(3)
            remainder = stripped[m.end():].strip()
            current = [remainder] if remainder else []
        else:
            current.append(stripped)

    if current:
        blocks.append(current)

    for block_lines in blocks:
        q = _parse_block(block_lines)
        if q:
            questions.append(q)

    return questions


def _parse_block(lines: list[str]) -> dict | None:
    """
    Parse a list of lines (one question block) into a dict.
    Returns None if the block cannot be parsed into a valid MCQ.
    """
    # Option: "A. text"  "A) text"  "A - text"  "(A) text"
    option_re = re.compile(
        r'^\s*[(\[]?\s*([A-Da-d])\s*[.):\-]\s*(.+)', re.IGNORECASE
    )
    # Answer: "✅ Correct Answer: C"  "Correct Answer: C"  "Answer: C"
    #         also inline: "Correct Answer: B Rationale: ..."
    answer_re = re.compile(
        r'^[^\w]*(?:Correct\s+)?Answer\s*[:\-]\s*([A-Da-d])',
        re.IGNORECASE,
    )
    # Short answer: "✅ B — explanation"  "✅ B explanation"
    short_answer_re = re.compile(
        r'^[^\w]*([A-Da-d])\s*(?:—|-|–|:)?\s*(.+)', re.IGNORECASE
    )
    # Rationale / Explanation  (may appear inline after answer on same line)
    rationale_inline_re = re.compile(
        r'(?:Rationale|Explanation|Discussion)\s*[:\-]\s*(.*)',
        re.IGNORECASE,
    )
    rationale_start_re = re.compile(
        r'^\s*(?:Rationale|Explanation|Discussion)\s*[:\-]\s*(.*)',
        re.IGNORECASE,
    )

    question_lines: list[str] = []
    options: dict[str, str] = {}
    correct: str | None = None
    explanation_lines: list[str] = []
    mode = "question"

    for line in lines:
        if not line:
            continue

        opt_m = option_re.match(line)
        ans_m = answer_re.match(line)
        rat_m = rationale_start_re.match(line)

        # ── Answer line ──────────────────────────────────────────────────────
        if ans_m and mode in ("options", "answer", "explanation"):
            correct = ans_m.group(1).lower()
            # check for inline rationale on the same line: "Correct Answer: B Rationale: ..."
            inline = rationale_inline_re.search(line)
            if inline:
                explanation_lines.append(inline.group(1))
            mode = "answer"

        # ── Rationale / Explanation line ─────────────────────────────────────
        elif rat_m:
            explanation_lines.append(rat_m.group(1))
            mode = "explanation"

        elif mode == "explanation":
            explanation_lines.append(line)

        # ── Option line ──────────────────────────────────────────────────────
        elif opt_m and (mode in ("question", "options") or not options):
            letter = opt_m.group(1).lower()
            options[letter] = opt_m.group(2).strip()
            mode = "options"

        # ── Post-options line that looks like "✅ B — explanation" ────────────
        elif mode == "options" and len(options) == 4:
            # Could be the short-answer format: ✅ B — explanation
            sa_m = short_answer_re.match(line)
            if sa_m and sa_m.group(1).lower() in "abcd":
                correct = sa_m.group(1).lower()
                explanation_lines.append(sa_m.group(2).strip())
                mode = "answer"
            else:
                # continuation of last option
                last_key = list(options.keys())[-1]
                options[last_key] += " " + line

        # ── Question text ─────────────────────────────────────────────────────
        elif mode == "question":
            question_lines.append(line)

        elif mode == "options" and options:
            last_key = list(options.keys())[-1]
            options[last_key] += " " + line

    if not question_lines or len(options) < 4 or correct is None or correct not in "abcd":
        return None

    return {
        "text": " ".join(question_lines),
        "option_a": options.get("a", ""),
        "option_b": options.get("b", ""),
        "option_c": options.get("c", ""),
        "option_d": options.get("d", ""),
        "correct_option": correct,
        "explanation": " ".join(explanation_lines),
    }


# ---------------------------------------------------------------------------
# Command
# ---------------------------------------------------------------------------

class Command(BaseCommand):
    help = "Seed the database with the CPHQ course, chapters, topics, and questions."

    def add_arguments(self, parser):
        parser.add_argument(
            "--flush",
            action="store_true",
            help="Delete existing CPHQ course data before loading.",
        )
        parser.add_argument(
            "--course-dir",
            default=None,
            help="Path to the 'CPHQ Course' folder (default: <BASE_DIR>/CPHQ Course).",
        )

    def handle(self, *args, **options):
        base_dir: Path = Path(settings.BASE_DIR)
        course_dir: Path = Path(options["course_dir"]) if options["course_dir"] else base_dir / "CPHQ Course"

        if not course_dir.exists():
            raise CommandError(f"Course folder not found: {course_dir}")

        if options["flush"]:
            self._flush()

        self.stdout.write(self.style.MIGRATE_HEADING("=== Loading CPHQ data ==="))

        # ── Category & Course ──────────────────────────────────────────────
        category, _ = Category.objects.get_or_create(
            slug="healthcare-quality",
            defaults={"name": "Healthcare Quality", "description": "CPHQ exam preparation materials."},
        )

        course, created = Course.objects.get_or_create(
            slug="cphq-exam-preparation",
            defaults={
                "title": "CPHQ Exam Preparation",
                "category": category,
                "about": (
                    "Comprehensive preparation for the Certified Professional in Healthcare Quality (CPHQ) exam. "
                    "Covers all exam domains with video lectures, quizzes, and mock examinations."
                ),
                "description": (
                    "This course prepares healthcare professionals for the CPHQ certification exam. "
                    "It includes in-depth video instruction, domain-specific quizzes, and full-length mock exams "
                    "aligned with the NAHQ exam blueprint."
                ),
                "is_published": True,
                "order": 1,
            },
        )
        self._log("created" if created else "already exists", "Course", course.title)

        # ── Domains → Chapters + Topics ────────────────────────────────────
        domains = [
            {
                "order": 1,
                "folder": course_dir / "Domain 1",
                "title": "Domain 1: Patient Safety",
                "topic_name": "Patient Safety",
                "video_glob": "Domain 1/Chapter 1/videos",
                "quiz_doc": course_dir / "Domain 1" / "Chapter 1" / "quizes" / "Quiz.docx",
                "mock_doc": course_dir / "Domain 1" / "Domain 1 mock exam" / "domain 1 - Mock exam.docx",
            },
            {
                "order": 2,
                "folder": course_dir / "Domain 2",
                "title": "Domain 2: Quality Review & Accountability",
                "topic_name": "Quality Review & Accountability",
                "video_glob": "Domain 2/Chapter 1/videos",
                "quiz_doc": course_dir / "Domain 2" / "Chapter 1" / "quizes" / "quiz.docx",
                "mock_doc": course_dir / "Domain 2" / "Moc exam question" / "Domain 2 - Mock exam.docx",
            },
        ]

        media_video_dir = Path(settings.MEDIA_ROOT) / "courses" / "videos"
        media_video_dir.mkdir(parents=True, exist_ok=True)

        for domain in domains:
            chapter, topic = self._load_domain(course, domain, course_dir, media_video_dir)

        # ── Global 140-question mock exam ──────────────────────────────────
        global_mock = course_dir / "CPHQ Mock Exam - 140 Questions.docx"
        if global_mock.exists():
            self._load_questions(
                global_mock,
                topic=None,           # no specific topic → general
                question_type="mock",
                label="Global mock exam (140 Qs)",
            )
        else:
            self.stdout.write(self.style.WARNING(f"  Not found: {global_mock.name}"))

        self.stdout.write(self.style.SUCCESS("\n=== Done! CPHQ data loaded successfully. ==="))

    # ── Domain loader ──────────────────────────────────────────────────────

    def _load_domain(self, course: Course, domain: dict, course_dir: Path, media_video_dir: Path):
        # Topic
        topic, created = Topic.objects.get_or_create(
            name=domain["topic_name"],
            defaults={"description": f"Questions covering the {domain['topic_name']} domain of the CPHQ exam."},
        )
        self._log("created" if created else "already exists", "Topic", topic.name)

        # Chapter
        chapter, created = Chapter.objects.get_or_create(
            course=course,
            order=domain["order"],
            defaults={
                "title": domain["title"],
                "is_published": True,
                "duration_minutes": 0,
            },
        )
        if not created:
            chapter.title = domain["title"]
            chapter.is_published = True
            chapter.save()
        self._log("created" if created else "updated", "Chapter", chapter.title)

        # Video
        video_dir = course_dir / domain["video_glob"]
        if video_dir.exists():
            mp4_files = list(video_dir.glob("*.mp4"))
            if mp4_files:
                src = mp4_files[0]
                dest = media_video_dir / src.name
                if not dest.exists():
                    self.stdout.write(f"  Copying video: {src.name} (~{src.stat().st_size // (1024*1024)} MB) ...")
                    shutil.copy2(src, dest)
                    self.stdout.write(self.style.SUCCESS("  Video copied."))
                else:
                    self.stdout.write(f"  Video already in media: {src.name}")

                rel_path = f"courses/videos/{dest.name}"
                if chapter.video_file.name != rel_path:
                    chapter.video_file.name = rel_path
                    # Estimate duration from filename or leave as 0
                    chapter.save()
            else:
                self.stdout.write(self.style.WARNING(f"  No .mp4 found in {video_dir}"))
        else:
            self.stdout.write(self.style.WARNING(f"  Video folder not found: {video_dir}"))

        # Quiz questions
        if domain["quiz_doc"].exists():
            self._load_questions(domain["quiz_doc"], topic, "quiz", f"{domain['topic_name']} quiz")
        else:
            self.stdout.write(self.style.WARNING(f"  Quiz doc not found: {domain['quiz_doc'].name}"))

        # Mock exam questions
        if domain["mock_doc"].exists():
            self._load_questions(domain["mock_doc"], topic, "mock", f"{domain['topic_name']} mock exam")
        else:
            self.stdout.write(self.style.WARNING(f"  Mock doc not found: {domain['mock_doc'].name}"))

        return chapter, topic

    # ── Question loader ────────────────────────────────────────────────────

    def _load_questions(self, docx_path: Path, topic, question_type: str, label: str):
        self.stdout.write(f"\n  Parsing: {docx_path.name}")
        try:
            text = _extract_text_from_docx(docx_path)
        except Exception as exc:
            self.stdout.write(self.style.ERROR(f"  Failed to read {docx_path.name}: {exc}"))
            return

        parsed = _parse_questions(text)
        self.stdout.write(f"  Parsed {len(parsed)} questions from {label}")

        created_count = 0
        skipped_count = 0
        for q in parsed:
            if not q["text"] or not q["option_a"]:
                skipped_count += 1
                continue
            # Avoid duplicates by matching question text
            if Question.objects.filter(text=q["text"]).exists():
                skipped_count += 1
                continue
            Question.objects.create(
                topic=topic,
                question_type=question_type,
                difficulty="medium",
                text=q["text"],
                option_a=q["option_a"],
                option_b=q["option_b"],
                option_c=q["option_c"],
                option_d=q["option_d"],
                correct_option=q["correct_option"],
                explanation=q["explanation"],
                is_active=True,
            )
            created_count += 1

        self.stdout.write(
            self.style.SUCCESS(f"  Created {created_count} questions") +
            (f" (skipped {skipped_count})" if skipped_count else "")
        )

    # ── Flush helper ───────────────────────────────────────────────────────

    def _flush(self):
        self.stdout.write(self.style.WARNING("Flushing existing CPHQ data..."))
        Course.objects.filter(slug="cphq-exam-preparation").delete()
        Topic.objects.filter(name__in=["Patient Safety", "Quality Review & Accountability"]).delete()
        Question.objects.all().delete()
        self.stdout.write(self.style.WARNING("Existing data removed.\n"))

    def _log(self, action: str, model: str, name: str):
        colour = self.style.SUCCESS if action == "created" else self.style.HTTP_INFO
        self.stdout.write(colour(f"  [{action}] {model}: {name}"))
